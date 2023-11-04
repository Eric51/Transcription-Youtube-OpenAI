from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QTextEdit, QLineEdit, QLabel, QComboBox
from PyQt5.QtCore import QThread, pyqtSignal
from docx import Document
import yt_dlp
import sys
import openai
import tempfile
import os
import json
from pydub import AudioSegment
import re

def load_config():
    if os.path.exists('config.json'):
        with open('config.json', 'r') as f:
            return json.load(f)
    return {}

def save_config(config):
    with open('config.json', 'w') as f:
        json.dump(config, f)

def pose_question(self, texte):
    # Charger la configuration depuis le fichier JSON
    config = load_config()
    question = config.get("question", "Question par défaut si non spécifiée.")
    openai.api_key = config.get("api_key", "")
    model = config.get("model", "gpt-4")
    self.segment_duration_seconds = int(config.get("segmentation", 60))
    
    # Créer la liste de messages pour la conversation
    messages_a_envoyer = [
        {"role": "user", "content": question},
        {"role": "assistant", "content": texte}
        
    ]
    
    # Poser la question à OpenAI pour obtenir une réponse
    response = openai.ChatCompletion.create(
        model=model,
        temperature=0.2,
        messages=messages_a_envoyer
    )
    
    # Extraire la réponse de la dernière sortie
    answer = response['choices'][0]['message']['content']
    
    return answer  
        
        
class TranscriptionThread(QThread):
    textReady = pyqtSignal(str)

    def __init__(self):
        super().__init__()
        self.title = ""
        self.url = ""
        self.segment_duration_seconds = 90 

    def run(self):
        config = load_config()
        openai.api_key = config.get("api_key", "")
        model = config.get("model", "whisper-1")

        url = self.url
        options = {
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'wav',
            }],
            'ffmpeg_location': 'C:\\ffmpeg\\bin',
            'outtmpl': 'audio',
            'quiet': True
        }

        answers = []  # Liste pour stocker les réponses

        with yt_dlp.YoutubeDL(options) as ydl:
            info_dict = ydl.extract_info(url, download=False)
            self.title = info_dict.get('title', None)
            ydl.download([url])

        doc = Document()
        if self.title:
            doc.add_heading(self.title, 0)
        else:
            doc.add_heading('Transcription YouTube', 0)

        audio = AudioSegment.from_wav("audio.wav")
        segments = [audio[i:i+self.segment_duration_seconds*1000] for i in range(0, len(audio), self.segment_duration_seconds*1000)]  

        current_time_seconds = 0

        for idx, segment in enumerate(segments):
            minutes, seconds = divmod(current_time_seconds, 60)
            timestamp = f"{minutes:02d}:{seconds:02d}"

            temp_filename = tempfile.mktemp(".wav")
            segment.export(temp_filename, format="wav")

            with open(temp_filename, "rb") as audio_file:
                transcript = openai.Audio.transcribe("whisper-1", audio_file)

            text = transcript.get("text", "Texte non disponible.")
            text = text.replace('. ', '.\n')  # Ajout d'un retour à la ligne après chaque phrase
            
            answer = pose_question(self, text)  # Obtention de la réponse
            
            answers.append(f"[{timestamp}] {answer}\n")  # Ajout de la réponse à la liste

            doc.add_paragraph(f"[{timestamp}] {text}\n")
            self.textReady.emit(f"[{timestamp}] {text}\n")

            current_time_seconds += self.segment_duration_seconds 

            os.remove(temp_filename)

        # Ajoutez les réponses à la fin du document
        for answer in answers:
            doc.add_paragraph(answer)

        if self.title:
            sanitized_title = re.sub(r'[\\/*?:"<>|]', '', self.title)
            doc.save(f"{sanitized_title}.docx")
            self.textReady.emit(f"Transcription sauvegardée dans '{self.title}.docx'.")
        else:
            doc.save('Transcription.docx')
            self.textReady.emit("Transcription sauvegardée dans 'Transcription.docx'.")

class ConfigWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Configuration')
        layout = QVBoxLayout()

        config = load_config()
        api_key = config.get("api_key", "")
        model = config.get("model", "")
        question = config.get("question", "")
        
        self.apiKeyInput = QLineEdit(self)
        self.apiKeyInput.setPlaceholderText('Entrez votre clé API OpenAI ici...')
        self.apiKeyInput.setText(api_key)

        self.modelInput = QComboBox(self)
        self.modelInput.addItem("gpt-4")
        self.modelInput.addItem("gpt-3.5-turbo")
        self.modelInput.addItem("gpt-3.5-turbo-16k")
        index = self.modelInput.findText(model)
        if index >= 0:
            self.modelInput.setCurrentIndex(index)

        self.questionInput = QLineEdit(self)
        self.questionInput.setPlaceholderText('Entrez votre question ici...')
        self.questionInput.setText(question)

        layout.addWidget(QLabel('Clé API:'))
        layout.addWidget(self.apiKeyInput)
        layout.addWidget(QLabel('Modèle:'))
        layout.addWidget(self.modelInput)
        layout.addWidget(QLabel('Question:'))
        layout.addWidget(self.questionInput)
        
        self.segmentationInput = QLineEdit(self)
        self.segmentationInput.setPlaceholderText('Entrez la durée de la segmentation en secondes...')
        segmentation = config.get("segmentation", "60")
        self.segmentationInput.setText(str(segmentation))
        
        layout.addWidget(QLabel('Segmentation:'))
        layout.addWidget(self.segmentationInput)

        self.saveButton = QPushButton('Sauvegarder', self)
        self.saveButton.clicked.connect(self.save)
        layout.addWidget(self.saveButton)

        self.setLayout(layout)

    def save(self):
        api_key = self.apiKeyInput.text()
        model = self.modelInput.currentText()
        question = self.questionInput.text()
        segmentation = int(self.segmentationInput.text())
        save_config({"api_key": api_key, "model": model, "question": question, "segmentation": segmentation})
        self.close()

class App(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Transcripteur YouTube')
        layout = QVBoxLayout()

        self.urlInput = QLineEdit(self)
        self.urlInput.setPlaceholderText('Entrez l\'URL de la vidéo YouTube ici...')
        layout.addWidget(self.urlInput)

        self.transcribeButton = QPushButton('Transcrire', self)
        self.transcribeButton.clicked.connect(self.transcribe)
        layout.addWidget(self.transcribeButton)

        self.configButton = QPushButton('Configuration', self)
        self.configButton.clicked.connect(self.show_config)
        layout.addWidget(self.configButton)

        self.outputArea = QTextEdit(self)
        self.outputArea.setReadOnly(True)
        layout.addWidget(self.outputArea)

        self.setLayout(layout)
        self.show()

        self.transcriptionThread = TranscriptionThread()
        self.transcriptionThread.textReady.connect(self.updateText)

    def transcribe(self):
        url = self.urlInput.text()
        if not url:
            self.outputArea.setText('Veuillez entrer une URL.')
            return

        self.transcriptionThread.url = url
        self.transcriptionThread.start()
        

    def updateText(self, text):
        self.outputArea.append(text)

    def show_config(self):
        self.configWindow = ConfigWindow()
        self.configWindow.show()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())
